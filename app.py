from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from datetime import datetime, timedelta
import os
import uuid
import threading
import webbrowser

app = Flask(__name__)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

TEMPLATE_PATH = os.path.join(BASE_DIR, "template.docx")
OUTPUT_FOLDER = os.path.join(BASE_DIR, "generated")
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/generate", methods=["POST"])
def generate():
    name = request.form.get("name")
    reg = request.form.get("reg")
    utr = request.form.get("utr")
    payment = request.form.get("payment")
    plan = request.form.get("plan")
    start = request.form.get("start")
    duration = int(request.form.get("duration"))

    start_date = datetime.strptime(start, "%Y-%m-%d")
    end_date = start_date + timedelta(days=duration)

    start_formatted = start_date.strftime("%d-%m-%Y")
    end_formatted = end_date.strftime("%d-%m-%Y")

    doc = Document(TEMPLATE_PATH)

    replacements = {
        "{Name}": name,
        "{R}": reg,
        "{U}": utr,
        "{P}": payment,
        "{DP}": plan,
        "{SD}": start_formatted,
        "{ED}": end_formatted,
        "{TP}": str(duration)
    }

    def apply_style(run):
        run.font.name = "Aptos"
        run.font.size = Pt(16)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')

    def replace_all(doc, replacements):
        for para in doc.paragraphs:
            for key, value in replacements.items():
                if key in para.text:
                    for run in para.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)
                            apply_style(run)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in para.text:
                                for run in para.runs:
                                    if key in run.text:
                                        run.text = run.text.replace(key, value)
                                        apply_style(run)

    replace_all(doc, replacements)

    # Save DOCX
    unique_id = uuid.uuid4().hex
    docx_path = os.path.join(OUTPUT_FOLDER, f"{unique_id}.docx")

    doc.save(docx_path)

    return send_file(
        docx_path,
        as_attachment=True,
        download_name=f"Receipt {name}.docx"
    )


# 🔥 AUTO OPEN BROWSER
def open_browser():
    chrome_path = "C:/Program Files/Google/Chrome/Application/chrome.exe %s"
    try:
        webbrowser.get(chrome_path).open("http://127.0.0.1:5000")
    except:
        webbrowser.open("http://127.0.0.1:5000")


if __name__ == "__main__":
    threading.Timer(1.5, open_browser).start()
    app.run(debug=False)